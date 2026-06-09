"""IGキャンペーン候補抽出ツール 操作マニュアル Word生成スクリプト"""

import io
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── helpers ──────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """セルの背景色を設定"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None, header_color="1B4F72"):
    """ヘッダー付きのスタイル済みテーブルを追加"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # ヘッダー行
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # データ行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "EBF5FB")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_flowchart_table(doc, steps, colors=None):
    """フローチャート風の図をテーブルで描画"""
    if colors is None:
        colors = ["2E86C1"] * len(steps)

    for i, (title, detail) in enumerate(steps):
        # ボックス
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        cell.width = Cm(14)
        set_cell_shading(cell, colors[i])

        # タイトル
        p_title = cell.paragraphs[0]
        run_t = p_title.add_run(title)
        run_t.bold = True
        run_t.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run_t.font.size = Pt(11)
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 詳細
        if detail:
            p_detail = cell.add_paragraph()
            run_d = p_detail.add_run(detail)
            run_d.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run_d.font.size = Pt(9)
            p_detail.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 矢印（最後以外）
        if i < len(steps) - 1:
            p_arrow = doc.add_paragraph()
            p_arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_a = p_arrow.add_run("▼")
            run_a.font.size = Pt(18)
            run_a.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
            p_arrow.paragraph_format.space_before = Pt(2)
            p_arrow.paragraph_format.space_after = Pt(2)


def add_info_box(doc, text, bg_color="FEF9E7", border_color="F39C12"):
    """注意・情報ボックスを追加"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(15)
    set_cell_shading(cell, bg_color)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x7D, 0x6A, 0x00)

    # 左ボーダー
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:color="{border_color}"/>'
        f"</w:tcBorders>"
    )
    tc_pr.append(borders)


def add_heading_with_number(doc, number, text, level=1):
    """番号付き見出しを追加"""
    h = doc.add_heading(level=level)
    run_num = h.add_run(f"{number}. ")
    run_num.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    run_text = h.add_run(text)
    if level == 1:
        run_text.font.size = Pt(18)
    return h


def add_step_paragraph(doc, step_num, text, bold_part=None):
    """ステップ番号付きパラグラフ"""
    p = doc.add_paragraph()
    run_num = p.add_run(f"  {step_num}  ")
    run_num.bold = True
    run_num.font.size = Pt(11)
    run_num.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 番号の背景 (shading on run is tricky, use text approach)
    if bold_part:
        run_b = p.add_run(f" {bold_part}")
        run_b.bold = True
        run_b.font.size = Pt(11)
        remaining = text.replace(bold_part, "", 1)
        if remaining:
            run_r = p.add_run(remaining)
            run_r.font.size = Pt(11)
    else:
        run_t = p.add_run(f" {text}")
        run_t.font.size = Pt(11)
    return p


# ── メインドキュメント生成 ──────────────────────────────

doc = Document()

# ページ設定
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

# デフォルトフォント
style = doc.styles["Normal"]
font = style.font
font.name = "Yu Gothic"
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")

# ━━━━━━━━ 表紙 ━━━━━━━━
doc.add_paragraph()  # 空行
doc.add_paragraph()

# タイトルボックス
tbl_cover = doc.add_table(rows=1, cols=1)
tbl_cover.alignment = WD_TABLE_ALIGNMENT.CENTER
cell_cover = tbl_cover.rows[0].cells[0]
set_cell_shading(cell_cover, "1B4F72")
cell_cover.width = Cm(16)

p_cover1 = cell_cover.paragraphs[0]
p_cover1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = p_cover1.add_run("\n")
r1.font.size = Pt(8)

p_cover2 = cell_cover.add_paragraph()
p_cover2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_cover2.add_run("IGキャンペーン候補抽出ツール")
r2.bold = True
r2.font.size = Pt(28)
r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

p_cover3 = cell_cover.add_paragraph()
p_cover3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p_cover3.add_run("操作マニュアル")
r3.font.size = Pt(20)
r3.font.color.rgb = RGBColor(0x85, 0xC1, 0xE9)

p_cover4 = cell_cover.add_paragraph()
p_cover4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p_cover4.add_run("\n")
r4.font.size = Pt(8)

doc.add_paragraph()
p_ver = doc.add_paragraph()
p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_ver = p_ver.add_run("Version 1.0 ｜ 2026年4月")
r_ver.font.size = Pt(12)
r_ver.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

doc.add_page_break()

# ━━━━━━━━ 目次 ━━━━━━━━
h_toc = doc.add_heading("目次", level=1)
h_toc.runs[0].font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

toc_items = [
    ("1", "ツールの概要"),
    ("2", "アクセス方法"),
    ("3", "初回セットアップ：アカウント登録"),
    ("4", "キャンペーン当選者を選定する手順"),
    ("5", "ジョブの確認・CSVダウンロード"),
    ("6", "CSVの見方と当選者選定の進め方"),
    ("7", "アカウント管理（セッションID更新・削除）"),
    ("8", "よくあるエラーと対処法"),
    ("9", "運用フロー全体図"),
]

for num, title in toc_items:
    p = doc.add_paragraph()
    run_n = p.add_run(f"  {num}  ")
    run_n.font.size = Pt(12)
    run_n.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    run_n.bold = True
    run_t = p.add_run(f"  {title}")
    run_t.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(4)

# 目次の下に全体概要図
doc.add_paragraph()
p_overview_label = doc.add_paragraph()
p_overview_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_ol = p_overview_label.add_run("【 全体の流れ 】")
r_ol.bold = True
r_ol.font.size = Pt(12)
r_ol.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

add_flowchart_table(doc, [
    ("STEP 1：事前準備", "セッションID取得 → アカウント登録"),
    ("STEP 2：抽出実行", "投稿URL入力 → モード選択 → 抽出開始"),
    ("STEP 3：結果取得", "ジョブ完了確認 → CSVダウンロード"),
    ("STEP 4：当選者選定", "候補絞り込み → プロフィール確認 → 当選連絡"),
], colors=["1B4F72", "2E86C1", "2ECC71", "E67E22"])

doc.add_page_break()

# ━━━━━━━━ 1. ツールの概要 ━━━━━━━━
add_heading_with_number(doc, "1", "ツールの概要")

p_desc = doc.add_paragraph(
    "Instagramのキャンペーン投稿に対して、いいね・コメントしたユーザーを自動抽出し、"
    "フォロワー数でフィルタリングした候補者リストをCSVで出力するツールです。"
)
p_desc.paragraph_format.space_after = Pt(12)

# 機能一覧を図解風テーブルで
h_func = doc.add_heading("できること", level=2)
features = [
    ("📥", "いいね取得", "指定した投稿の「いいね」したユーザー一覧を取得"),
    ("💬", "コメント取得", "指定した投稿の「コメント」したユーザー一覧を取得"),
    ("👥", "フォロワー数取得", "各ユーザーのフォロワー数を自動取得"),
    ("🔍", "フィルタリング", "最低フォロワー数で候補を絞り込み（例：1,000人未満を除外）"),
    ("📊", "CSV出力", "結果をCSVファイルとしてダウンロード"),
]
tbl_feat = doc.add_table(rows=len(features), cols=3)
tbl_feat.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_feat.style = "Table Grid"
for i, (icon, title, desc) in enumerate(features):
    c0 = tbl_feat.rows[i].cells[0]
    c0.width = Cm(1.5)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(icon)
    r0.font.size = Pt(16)
    set_cell_shading(c0, "EBF5FB")

    c1 = tbl_feat.rows[i].cells[1]
    c1.width = Cm(3.5)
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.size = Pt(10)

    c2 = tbl_feat.rows[i].cells[2]
    c2.width = Cm(11)
    p2 = c2.paragraphs[0]
    r2 = p2.add_run(desc)
    r2.font.size = Pt(10)

# 処理の流れ図
doc.add_paragraph()
p_flow_label = doc.add_paragraph()
p_flow_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_fl = p_flow_label.add_run("【 処理の流れ 】")
r_fl.bold = True
r_fl.font.size = Pt(11)
r_fl.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

flow_tbl = doc.add_table(rows=1, cols=9)
flow_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
flow_items = ["投稿URL\n入力", "→", "いいね/\nコメント\n取得", "→", "フォロワー数\n取得", "→", "フィルタ\nリング", "→", "CSV\n出力"]
flow_colors = ["2E86C1", "FFFFFF", "2E86C1", "FFFFFF", "2E86C1", "FFFFFF", "2E86C1", "FFFFFF", "2ECC71"]
for i, (item, color) in enumerate(zip(flow_items, flow_colors)):
    cell = flow_tbl.rows[0].cells[i]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(item)
    run.font.size = Pt(8)
    if color != "FFFFFF":
        set_cell_shading(cell, color)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.bold = True
    else:
        run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
        run.bold = True
        run.font.size = Pt(14)

doc.add_page_break()

# ━━━━━━━━ 2. アクセス方法 ━━━━━━━━
add_heading_with_number(doc, "2", "アクセス方法")

p = doc.add_paragraph("ブラウザで以下のURLにアクセスしてください（URLは管理者から共有されます）。")
p.paragraph_format.space_after = Pt(8)

# URL枠
tbl_url = doc.add_table(rows=1, cols=1)
tbl_url.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_url.style = "Table Grid"
cell_url = tbl_url.rows[0].cells[0]
set_cell_shading(cell_url, "F4F6F7")
p_url = cell_url.paragraphs[0]
p_url.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_url = p_url.add_run("https://（デプロイ先URL）")
r_url.font.name = "Consolas"
r_url.font.size = Pt(12)

doc.add_paragraph()
add_info_box(doc, "💡 ローカルで起動する場合：ターミナルで streamlit run app.py を実行し、表示されるURLにアクセス")

doc.add_page_break()

# ━━━━━━━━ 3. 初回セットアップ ━━━━━━━━
add_heading_with_number(doc, "3", "初回セットアップ：アカウント登録")

p = doc.add_paragraph("ツールを使うには、まずInstagramアカウントのセッションIDを登録する必要があります。")
p.paragraph_format.space_after = Pt(12)

# セットアップフロー図
p_setup_label = doc.add_paragraph()
p_setup_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sl = p_setup_label.add_run("【 セットアップの流れ 】")
r_sl.bold = True
r_sl.font.size = Pt(11)
r_sl.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

add_flowchart_table(doc, [
    ("Step 1：セッションID取得", "ChromeでInstagramにログイン → 開発者ツールからコピー"),
    ("Step 2：アカウント登録", "アカウント管理タブ → 表示名・ユーザー名・セッションIDを入力"),
    ("完了！", "抽出ツールが利用可能になります"),
], colors=["2E86C1", "2E86C1", "2ECC71"])

doc.add_paragraph()

# 3-1
h31 = doc.add_heading("3-1. セッションIDの取得手順（Chrome）", level=2)

steps_session = [
    ("①", "Chromeで instagram.com を開き、キャンペーン用アカウントでログイン"),
    ("②", "F12キー（Macは Cmd + Option + I）で開発者ツールを開く"),
    ("③", "上部メニューの「Application」タブをクリック"),
    ("④", "左メニューの「Cookies」→「https://www.instagram.com」をクリック"),
    ("⑤", "一覧から「sessionid」を探し、Value列の値をコピー"),
]
for num, text in steps_session:
    add_step_paragraph(doc, num, text)

doc.add_paragraph()
add_info_box(doc, "⚠️ セッションIDは機密情報です。Slack等で共有せず、ツール上で直接入力してください。",
             bg_color="FDEDEC", border_color="E74C3C")

doc.add_paragraph()

# 3-2
h32 = doc.add_heading("3-2. アカウントの登録", level=2)

p = doc.add_paragraph("画面上部の「⚙️ アカウント管理」タブを開き、「➕ 新規アカウント追加」セクションに以下を入力します。")
p.paragraph_format.space_after = Pt(8)

add_styled_table(doc,
    ["項目", "入力内容", "例"],
    [
        ["表示名（社内用）", "チーム内で識別しやすい名前", "クライアントA"],
        ["Instagramユーザー名", "@なしのユーザー名", "ig_account_name"],
        ["セッションID", "上記手順でコピーした値", "（長い英数字の文字列）"],
    ],
    col_widths=[4, 6, 5]
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.add_run("入力後、").font.size = Pt(10.5)
r = p.add_run("「追加する」ボタンをクリック")
r.bold = True
r.font.size = Pt(10.5)
p.add_run("して登録完了です。").font.size = Pt(10.5)

doc.add_page_break()

# ━━━━━━━━ 4. キャンペーン当選者を選定する手順 ━━━━━━━━
add_heading_with_number(doc, "4", "キャンペーン当選者を選定する手順")

# ステップ1
h_s1 = doc.add_heading("ステップ1：抽出条件を設定する", level=2)

p = doc.add_paragraph("「📊 抽出ツール」タブを開き、以下の項目を設定します。")
p.paragraph_format.space_after = Pt(8)

add_styled_table(doc,
    ["項目", "説明"],
    [
        ["使用するアカウント", "登録済みのアカウントから選択"],
        ["投稿URL", "キャンペーン対象の投稿URLを貼り付け\n（https://www.instagram.com/p/XXXXXXXXX/ の形式）"],
        ["抽出モード", "下記の3パターンから選択"],
        ["最低フォロワー数", "この数値未満のユーザーを除外（デフォルト：1,000）"],
    ],
    col_widths=[4.5, 11.5]
)

doc.add_paragraph()

# 抽出モード比較図
p_mode_label = doc.add_paragraph()
p_mode_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_ml = p_mode_label.add_run("【 抽出モードの選び方 】")
r_ml.bold = True
r_ml.font.size = Pt(11)
r_ml.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

# 3カラム比較テーブル
tbl_mode = doc.add_table(rows=3, cols=3)
tbl_mode.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_mode.style = "Table Grid"

modes = [
    ("いいね＋コメント\n（どちらかでOK）", "最も広い候補リスト", "2E86C1"),
    ("いいねのみ", "コメント不要の\nキャンペーン向け", "27AE60"),
    ("コメントのみ", "コメント内容も\nCSVに出力", "E67E22"),
]

# ヘッダー行
for i, (mode_name, _, color) in enumerate(modes):
    cell = tbl_mode.rows[0].cells[i]
    set_cell_shading(cell, color)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(mode_name)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# アイコン行
icons = ["❤️ + 💬", "❤️", "💬"]
for i, icon in enumerate(icons):
    cell = tbl_mode.rows[1].cells[i]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(icon)
    r.font.size = Pt(20)

# 説明行
for i, (_, desc, _) in enumerate(modes):
    cell = tbl_mode.rows[2].cells[i]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(desc)
    r.font.size = Pt(9)

doc.add_paragraph()

# ステップ2
h_s2 = doc.add_heading("ステップ2：抽出を実行する", level=2)

steps_exec = [
    ("①", "設定を確認して「抽出開始」ボタンをクリック"),
    ("②", "「ジョブを登録しました」と表示されたら成功"),
    ("③", "「📋 ジョブ一覧」タブに移動して進捗を確認"),
]
for num, text in steps_exec:
    add_step_paragraph(doc, num, text)

doc.add_paragraph()
add_info_box(doc, "⏱ 処理はバックグラウンドで実行されます。ブラウザを閉じずにお待ちください。\n"
                  "対象ユーザーが多い場合、フォロワー数取得に時間がかかります（数百人で10〜30分程度）。")

doc.add_page_break()

# ━━━━━━━━ 5. ジョブの確認・CSVダウンロード ━━━━━━━━
add_heading_with_number(doc, "5", "ジョブの確認・CSVダウンロード")

p = doc.add_paragraph("「📋 ジョブ一覧」タブを開き、ジョブのステータスを確認します。")
p.paragraph_format.space_after = Pt(12)

# ステータス図解
p_status_label = doc.add_paragraph()
p_status_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_stl = p_status_label.add_run("【 ステータスの流れ 】")
r_stl.bold = True
r_stl.font.size = Pt(11)
r_stl.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

# ステータスフロー（横並び）
flow_status = doc.add_table(rows=2, cols=7)
flow_status.alignment = WD_TABLE_ALIGNMENT.CENTER

statuses = [
    ("🕐 待機中", "BDC3C7"),
    ("→", "FFFFFF"),
    ("⏳ 実行中", "F39C12"),
    ("→", "FFFFFF"),
    ("✅ 完了", "27AE60"),
    ("", "FFFFFF"),
    ("❌ エラー", "E74C3C"),
]
descs = ["処理待ち", "", "処理中", "", "CSVダウンロード可", "", "エラー内容を確認"]

for i, ((label, color), desc) in enumerate(zip(statuses, descs)):
    cell_top = flow_status.rows[0].cells[i]
    p = cell_top.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(10)
    if color not in ("FFFFFF",):
        set_cell_shading(cell_top, color)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    else:
        r.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
        r.font.size = Pt(14)

    cell_bot = flow_status.rows[1].cells[i]
    p2 = cell_bot.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(desc)
    r2.font.size = Pt(8)
    r2.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

doc.add_paragraph()

steps_dl = [
    ("①", "「📋 ジョブ一覧」タブを開く"),
    ("②", "ステータスが ✅完了 になるまで待つ"),
    ("③", "「📥 CSV」ボタンをクリックしてダウンロード"),
    ("④", "不要なジョブは「削除」ボタンで削除可能"),
]
for num, text in steps_dl:
    add_step_paragraph(doc, num, text)

doc.add_page_break()

# ━━━━━━━━ 6. CSVの見方 ━━━━━━━━
add_heading_with_number(doc, "6", "CSVの見方と当選者選定の進め方")

h_csv = doc.add_heading("CSVの列", level=2)

add_styled_table(doc,
    ["列名", "内容"],
    [
        ["id_name", "ユーザー名（@付き）"],
        ["profile_url", "Instagramプロフィールページへのリンク"],
        ["follower_count", "フォロワー数"],
        ["comment", "コメント内容（コメント抽出モード時のみ）"],
    ],
    col_widths=[4, 12]
)

doc.add_paragraph()
add_info_box(doc, "📋 CSVはフォロワー数の多い順にソートされています。")

doc.add_paragraph()

# 当選者選定フロー
h_flow = doc.add_heading("当選者選定の推奨フロー", level=2)

add_flowchart_table(doc, [
    ("① CSVダウンロード", "スプレッドシートで開く"),
    ("② 応募条件の確認", "フォロー・いいね・コメントなどキャンペーン参加条件を確認"),
    ("③ 候補者の絞り込み", "フォロワー数・コメント内容をもとに候補を選定"),
    ("④ プロフィール確認", "profile_url からアカウントの質を目視確認"),
    ("⑤ 当選者リスト確定", "チーム内で最終確認し、当選者を決定"),
    ("⑥ 当選連絡", "DMなどで当選者へ連絡"),
], colors=["2E86C1", "3498DB", "2980B9", "1F618D", "1B4F72", "2ECC71"])

doc.add_page_break()

# ━━━━━━━━ 7. アカウント管理 ━━━━━━━━
add_heading_with_number(doc, "7", "アカウント管理（セッションID更新・削除）")

# セッションID更新
h_update = doc.add_heading("セッションIDの更新", level=2)

add_info_box(doc, "⚠️ セッションIDには有効期限があります（数ヶ月）。「セッション切れ」エラーが出たら更新が必要です。",
             bg_color="FDEDEC", border_color="E74C3C")

doc.add_paragraph()

steps_update = [
    ("①", "「⚙️ アカウント管理」タブを開く"),
    ("②", "「🔑 セッションID更新」セクションで対象アカウントを選択"),
    ("③", "新しいセッションIDを入力して「更新する」をクリック"),
]
for num, text in steps_update:
    add_step_paragraph(doc, num, text)

doc.add_paragraph()

# 更新フロー
p_upd_label = doc.add_paragraph()
p_upd_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_ul = p_upd_label.add_run("【 セッションID更新の流れ 】")
r_ul.bold = True
r_ul.font.size = Pt(11)
r_ul.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

add_flowchart_table(doc, [
    ("セッション切れエラー発生", ""),
    ("Chromeで再ログイン", "開発者ツールからセッションIDをコピー"),
    ("アカウント管理タブで更新", "新しいセッションIDを入力 → 「更新する」"),
], colors=["E74C3C", "F39C12", "2ECC71"])

doc.add_paragraph()

# アカウント削除
h_del = doc.add_heading("アカウントの削除", level=2)
p = doc.add_paragraph("アカウント一覧から対象アカウントの「削除」ボタンをクリックします。")

doc.add_page_break()

# ━━━━━━━━ 8. よくあるエラーと対処法 ━━━━━━━━
add_heading_with_number(doc, "8", "よくあるエラーと対処法")

add_styled_table(doc,
    ["エラー", "原因", "対処法"],
    [
        ["セッション切れ", "セッションIDの有効期限切れ", "「アカウント管理」タブから\nセッションIDを再取得・更新"],
        ["投稿URLエラー", "URLに /p/ が含まれていない", "正しい投稿URLを確認して再入力"],
        ["アカウント未登録", "アカウントが登録されていない", "「アカウント管理」タブから\nアカウントを追加"],
        ["フォロワー数が空欄", "API制限でフォロワー数を\n取得できなかった", "時間を空けて再実行"],
    ],
    col_widths=[4, 5, 7],
    header_color="C0392B"
)

doc.add_paragraph()

# トラブルシューティングフロー
p_ts_label = doc.add_paragraph()
p_ts_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_tsl = p_ts_label.add_run("【 トラブルシューティングフロー 】")
r_tsl.bold = True
r_tsl.font.size = Pt(11)
r_tsl.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

# 判断分岐テーブル
tbl_ts = doc.add_table(rows=5, cols=3)
tbl_ts.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_ts.style = "Table Grid"

# Row 0: 問題発生
cell = tbl_ts.rows[0].cells[0]
cell.merge(tbl_ts.rows[0].cells[2])
set_cell_shading(cell, "E74C3C")
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("エラーが発生した！")
r.bold = True
r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# Row 1: 矢印
cell = tbl_ts.rows[1].cells[0]
cell.merge(tbl_ts.rows[1].cells[2])
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("▼ エラーメッセージを確認")
r.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

# Row 2: 分岐ヘッダー
for i, (label, color) in enumerate([
    ("セッション切れ？", "F39C12"),
    ("URL不正？", "F39C12"),
    ("その他？", "F39C12"),
]):
    cell = tbl_ts.rows[2].cells[i]
    set_cell_shading(cell, color)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(9)

# Row 3: 矢印
for i in range(3):
    cell = tbl_ts.rows[3].cells[i]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("▼")
    r.font.color.rgb = RGBColor(0xF3, 0x9C, 0x12)

# Row 4: 対処法
solutions = [
    ("セッションID\nを再取得・更新", "27AE60"),
    ("投稿URLの形式\nを確認して再入力", "27AE60"),
    ("時間を空けて\n再実行", "27AE60"),
]
for i, (text, color) in enumerate(solutions):
    cell = tbl_ts.rows[4].cells[i]
    set_cell_shading(cell, color)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(9)
    r.bold = True

doc.add_page_break()

# ━━━━━━━━ 9. 運用フロー全体図 ━━━━━━━━
add_heading_with_number(doc, "9", "運用フロー全体図")

p = doc.add_paragraph("キャンペーン投稿から当選者選定までの全体の流れです。")
p.paragraph_format.space_after = Pt(16)

add_flowchart_table(doc, [
    ("事前準備（初回のみ）", "① セッションIDを取得\n② アカウント管理タブでアカウント登録"),
    ("キャンペーン投稿が公開されたら", "③ 投稿URLをコピー"),
    ("抽出ツールタブで実行", "④ アカウント選択・URL貼付・モード選択・フォロワー数設定\n⑤ 「抽出開始」をクリック"),
    ("ジョブ一覧タブで進捗確認", "⑥ ✅完了 になったらCSVダウンロード"),
    ("当選者選定", "⑦ CSVをスプレッドシートで開く\n⑧ 応募条件・プロフィールを確認して候補を絞り込む\n⑨ チームで最終確認 → 当選者確定\n⑩ 当選者にDMで連絡"),
], colors=["1B4F72", "2E86C1", "3498DB", "27AE60", "E67E22"])

doc.add_paragraph()

# ━━━━━━━━ 注意事項 ━━━━━━━━
h_note = doc.add_heading("注意事項", level=1)
h_note.runs[0].font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

warnings = [
    ("🔒", "セッションIDは機密情報です。", "Slack等で共有せず、ツール上で直接入力してください。"),
    ("⚠️", "Instagram APIの制限に注意", "短時間に大量のリクエストを送るとアカウントが一時制限される可能性があります。連続実行は避けてください。"),
    ("👥", "複数人での同時利用が可能です。", "ジョブはキュー管理されています。"),
]

for icon, title, desc in warnings:
    tbl_w = doc.add_table(rows=1, cols=2)
    tbl_w.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_w.style = "Table Grid"

    c0 = tbl_w.rows[0].cells[0]
    c0.width = Cm(1.5)
    set_cell_shading(c0, "FDEDEC")
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(icon)
    r0.font.size = Pt(18)

    c1 = tbl_w.rows[0].cells[1]
    c1.width = Cm(14.5)
    p1 = c1.paragraphs[0]
    r_title = p1.add_run(title)
    r_title.bold = True
    r_title.font.size = Pt(10)
    p1_desc = c1.add_paragraph()
    r_desc = p1_desc.add_run(desc)
    r_desc.font.size = Pt(9.5)
    r_desc.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    # 少し間隔
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ━━━━━━━━ 保存 ━━━━━━━━
output_path = "/Users/takedayuuki/Desktop/campaign/docs/IGキャンペーン候補抽出ツール_操作マニュアル.docx"
doc.save(output_path)
print(f"✅ Word文書を生成しました: {output_path}")
